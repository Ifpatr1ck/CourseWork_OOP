import docx
import functions.formatting
import functions.table

import os

import functions.table

paths = []
def search():
    # Ищет доки
    for file in os.listdir(os.getcwd()):
        if file.endswith('.docx'):
            paths.append(file)
def FileName():
    for file in os.listdir(os.getcwd()):
        if file.endswith('.docx'):
            return "" + str(file)

def names(path):
    text = 'Наименование документа: ' + path
    return text
def filecheck():
    try:
        # Ищет доки
        search()
        for path in paths:
            # отдает док в переменную
            doc = docx.Document(path)
            functions.formatting.indent_img(doc.paragraphs)  # написание в первой строке, о проверке отступау "Рисунок № -"
            print("Загрузка 6%")
            functions.formatting.check_size_font(177800, doc.paragraphs)  # проверка размера шрифта
            print("Загрузка 12%")
            functions.formatting.check_name_font('Times New Roman', doc.paragraphs)  # проверка имени шрифта
            print("Загрузка 18%")
            functions.formatting.check_heading(1.25, doc.paragraphs)  # проверка заголовков
            print("Загрузка 24%")
            functions.formatting.check_first_line_indent(1.25, doc.paragraphs)  # проверка отступа первого абзаца
            print("Загрузка 30%")
            functions.formatting.check_line_spacing(1.5, doc.paragraphs)  # проверка отступа межстрочного
            print("Загрузка 36%")
            functions.table.check_sources_merged(doc)  # библ список?
            print("Загрузка 42%")
            functions.table.check_pic_merged(doc)  # изображения
            print("Загрузка 48%")
            functions.table.check_table_merged(doc)  # таблицы
            print("Загрузка 54%")
            functions.table.check_block(doc)
            print("Загрузка 60%")
            functions.table.check_link_sources(doc)
            print("Загрузка 66%")
            functions.table.check_num_sources(doc)
            print("Загрузка 72%")
            functions.table.check_link_pic(doc)
            print("Загрузка 78%")
            functions.table.check_num_pic(doc)
            print("Загрузка 84%")
            functions.table.check_num_tables(doc)
            print("Загрузка 90%")
            functions.table.check_link_tables(doc)
            doc.save(path)
            paths.remove(path)

    except Exception as e:
        print('Ошибка!')

